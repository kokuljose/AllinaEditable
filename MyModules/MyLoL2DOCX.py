from docx import Document
from docx.shared import Inches
from PyPDF2 import PdfFileWriter, PdfFileReader
from pdf2image import convert_from_path
import os
OUTPUT_DIRECTORY = os.path.dirname(__file__)[:-10]+"/Output"
PROCESSING_DIRECTORY = os.path.dirname(__file__)[:-10]+"/Processing"
def LoL2DOCX(detailedTable,filename,IsRevised):
    if IsRevised:
        sec1Pages=2
    else:
        sec1Pages=1
    document = Document()
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = True
    document.sections[0].left_margin = Inches(0)
    document.sections[0].top_margin = Inches(0)
    document.sections[0].right_margin = Inches(0)
    document.sections[0].bottom_margin = Inches(0)
    pdf_file = PROCESSING_DIRECTORY + "/" + filename+".pdf"
    inputpdf = PdfFileReader(open(pdf_file, "rb"))
    for i in range(0, sec1Pages):
        output = PdfFileWriter()
        output.addPage(inputpdf.getPage(i))
        with open(OUTPUT_DIRECTORY + "/document-page%s.pdf" % i, "wb+") as outputStream:
            output.write(outputStream)
        page = convert_from_path(OUTPUT_DIRECTORY + "/document-page%s.pdf" % i, 500,
                                 poppler_path=r'D:\Users\Arun\poppler-0.68.0\bin')
        page[0].save(OUTPUT_DIRECTORY + '/output_{}.jpg'.format(i), 'JPEG')
        document.add_picture(OUTPUT_DIRECTORY + '/output_{}.jpg'.format(i), width=Inches(8.5))

    new_section = document.add_section()
    new_section.left_margin = Inches(1.0)
    new_section.right_margin = Inches(1.0)
    new_section.top_margin = Inches(1.0)
    new_section.bottom_margin = Inches(1.0)
    table = document.add_table(1, 10)
    rowNum = 0
    for row in detailedTable:
        row_cells = table.add_row().cells
        rowNum = rowNum + 1
        colNum = 0
        for col in row:
            colNum = colNum + 1
            row_cells[colNum].text = str(col)
            if rowNum < 5 or len(row) > 5:
                run = row_cells[colNum].paragraphs[0].runs[0]
                run.font.bold = True
            # print(str(row))

    document.add_page_break()
    new_section = document.add_section()
    new_section.left_margin = Inches(0)
    new_section.right_margin = Inches(0)
    new_section.top_margin = Inches(0)
    new_section.bottom_margin = Inches(0)
    for i in range(sec1Pages, inputpdf.numPages):
        output = PdfFileWriter()
        output.addPage(inputpdf.getPage(i))
        with open(OUTPUT_DIRECTORY+"/document-page%s.pdf" % i, "wb+") as outputStream:
            output.write(outputStream)
        page = convert_from_path(OUTPUT_DIRECTORY+"/document-page%s.pdf" % i, 500,poppler_path = r'D:\Users\Arun\poppler-0.68.0\bin')
        page[0].save(OUTPUT_DIRECTORY+'/output_{}.jpg'.format(i), 'JPEG')
        document.add_picture(OUTPUT_DIRECTORY + '/output_{}.jpg'.format(i), width=Inches(8.5))
    document.save(OUTPUT_DIRECTORY + '/'+filename+'Reconciled.docx')